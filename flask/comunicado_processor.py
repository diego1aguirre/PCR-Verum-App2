"""
Processes a formatted comunicado .docx and produces a plain-formatted version:
- Aptos font, 12pt
- No paragraph spacing (space_before/space_after = 0)
- Single line spacing
- Empty paragraph between each content paragraph
- Justified text

Copied verbatim from Comunicado-app/processor.py.
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Styles that act as section headers — no empty separator is inserted
# between them and the immediately-following sub-item paragraph.
_HEADER_STYLES = {'MetodologasyAnalistas', 'MetodologíasyAnalistas'}


def _para_style(para_el):
    pStyle = para_el.find('.//' + qn('w:pStyle'))
    return pStyle.get(qn('w:val')) if pStyle is not None else 'Normal'


def _para_text(para_el):
    return ''.join(t.text or '' for t in para_el.iter(qn('w:t')))


def _para_is_list_item(para_el):
    """True if the paragraph has a <w:numPr> (bulleted/numbered list item)."""
    return para_el.find('.//' + qn('w:numPr')) is not None


def _table_is_multi_para(tbl_el):
    """True if any cell in the table contains more than one non-empty paragraph."""
    for tc in tbl_el.iter(qn('w:tc')):
        paras = [p for p in tc.findall(qn('w:p'))]
        non_empty = [p for p in paras if _para_text(p).strip()]
        if len(non_empty) > 1:
            return True
    return False


def _extract_items(doc):
    """
    Walk the document body in document order and return a list of items.
    Each item is a dict:
      {'text': str, 'blank': bool, 'after_header': bool}

    'blank'        — True for empty separator paragraphs
    'after_header' — True when this item immediately follows a header-style paragraph;
                     the caller will suppress the empty separator before it.
    """
    body = doc.element.body
    items = []
    prev_was_header = False

    for child in body:
        tag = child.tag

        # ── Body paragraph ────────────────────────────────────────────────
        if tag == qn('w:p'):
            text = _para_text(child)
            stripped = text.strip()
            style = _para_style(child)

            if stripped:
                if _para_is_list_item(child):
                    stripped = '-\t' + stripped
                items.append({
                    'text': stripped,
                    'blank': False,
                    'after_header': prev_was_header,
                    'suppress_sep': False,
                })
                prev_was_header = style in _HEADER_STYLES
            else:
                items.append({'text': text, 'blank': True, 'after_header': False, 'suppress_sep': False})

        # ── Table ─────────────────────────────────────────────────────────
        elif tag == qn('w:tbl'):
            if _table_is_multi_para(child):
                # Analyst-style table: each column has multiple paragraphs.
                for tr in child.findall(qn('w:tr')):
                    for tc in tr.findall(qn('w:tc')):
                        cell_lines = [
                            _para_text(p)
                            for p in tc.findall(qn('w:p'))
                            if _para_text(p).strip()
                        ]
                        for idx, line in enumerate(cell_lines):
                            items.append({
                                'text': line,
                                'blank': False,
                                'after_header': False,
                                'suppress_sep': idx > 0,
                            })
                        items.append({'text': '', 'blank': True, 'after_header': False, 'suppress_sep': False})
            else:
                # Rating-style table: N rows × M columns, single paragraph per cell.
                for row_idx, tr in enumerate(child.findall(qn('w:tr'))):
                    cells = tr.findall(qn('w:tc'))
                    row_parts = []
                    for tc in cells:
                        p = tc.find(qn('w:p'))
                        row_parts.append(_para_text(p) if p is not None else '')
                    non_empty_parts = [part for part in row_parts if part.strip()]
                    row_text = '\t\t'.join(non_empty_parts).strip()
                    if row_text:
                        items.append({'text': row_text, 'blank': False, 'after_header': False, 'suppress_sep': row_idx > 0})
                items.append({'text': '', 'blank': True, 'after_header': False, 'suppress_sep': False})

            prev_was_header = False

    return items


def _make_plain_paragraph(doc, text):
    """
    Add a paragraph with plain formatting:
      - Aptos font, 12pt
      - Justified
      - No space before/after
      - Single line spacing (auto 240)
    """
    para = doc.add_paragraph()

    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)

    if text:
        run = para.add_run(text)
        run.font.name = 'Aptos'
        run.font.size = Pt(12)

        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Aptos')
        rFonts.set(qn('w:hAnsi'), 'Aptos')
        rFonts.set(qn('w:asciiTheme'), 'minorHAnsi')
        rFonts.set(qn('w:hAnsiTheme'), 'minorHAnsi')

    return para


def _build_output_filename(input_path):
    base = os.path.splitext(os.path.basename(input_path))[0]
    base = re.sub(r'_input$', '', base)
    return f'ComPrensa_{base}_plain.docx'


def process_comunicado(input_path: str, output_path: str) -> str:
    """
    Read a formatted comunicado .docx and write a plain-formatted version.
    Returns the suggested download filename.
    """
    input_doc = Document(input_path)
    items = _extract_items(input_doc)

    if not any(not item['blank'] for item in items):
        raise ValueError('No text content found in the uploaded document.')

    out_doc = Document()

    for p in out_doc.paragraphs:
        p._element.getparent().remove(p._element)

    out_doc.styles['Normal'].font.name = 'Aptos'
    out_doc.styles['Normal'].font.size = Pt(12)

    prev_was_content = False

    for item in items:
        if item['blank']:
            if prev_was_content or out_doc.paragraphs:
                _make_plain_paragraph(out_doc, item['text'])
            prev_was_content = False
        else:
            if prev_was_content and not item['after_header'] and not item.get('suppress_sep'):
                _make_plain_paragraph(out_doc, '')
            _make_plain_paragraph(out_doc, item['text'])
            prev_was_content = True

    out_doc.save(output_path)
    return _build_output_filename(input_path)
